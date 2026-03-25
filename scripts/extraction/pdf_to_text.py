import fitz  # PyMuPDF
import pytesseract
from docx import Document
from PIL import Image
from pathlib import Path

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from config import RAW_PDF_DIR, EXTRACTED_TEXT_DIR

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def extract_text_pymupdf(pdf_path):
    text = ""
    doc = fitz.open(pdf_path)
    for page in doc:
        text += page.get_text()
    doc.close()
    return text.strip()


def extract_text_ocr(pdf_path):
    text = ""
    doc = fitz.open(pdf_path)
    for page in doc:
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text += pytesseract.image_to_string(img)
    doc.close()
    return text.strip()


def extract_text_docx(docx_path):
    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs]

    # Also extract text from tables (exam papers often use tables for headers)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return "\n".join(paragraphs).strip()


def process_file(file_path, output_dir=None):
    """Extract text from a PDF or Word document. Returns path to the output .txt file."""
    file_path = Path(file_path)
    output_dir = Path(output_dir) if output_dir else EXTRACTED_TEXT_DIR
    ext = file_path.suffix.lower()

    output_txt = output_dir / file_path.with_suffix(".txt").name

    print(f"  [extract] Processing: {file_path.name}")

    if ext == ".docx":
        text = extract_text_docx(file_path)
    elif ext == ".pdf":
        text = extract_text_pymupdf(file_path)
        if len(text) < 200:
            print(f"  [extract] Low text detected, switching to OCR: {file_path.name}")
            text = extract_text_ocr(file_path)
    else:
        print(f"  [extract] Unsupported format: {ext}")
        return None

    output_txt.write_text(text, encoding="utf-8")
    print(f"  [extract] Saved: {output_txt.name}")

    return output_txt


# Keep backward compatibility
process_pdf = process_file


def process_folder(folder_path=None):
    """Process all PDFs and Word docs in a folder. Returns list of output .txt paths."""
    folder = Path(folder_path) if folder_path else RAW_PDF_DIR
    results = []

    for f in sorted(folder.rglob("*")):
        if f.suffix.lower() in SUPPORTED_EXTENSIONS:
            txt_path = process_file(f)
            if txt_path:
                results.append(txt_path)

    return results


if __name__ == "__main__":
    process_folder()
